import sys

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import pymysql
import ast


def getQuestions(id):
    sql = f"select description from tb_questions where id = {id}"
    cursor.execute(sql)
    return cursor.fetchall()[0]


def getOptions(id):
    sql = f"select options from tb_questions where id = {id}"
    cursor.execute(sql)
    options = ast.literal_eval(cursor.fetchall()[0][0])
    return options  # 返回获取的选项列表


def addSelectQuestion(selectQuestionlist):
    global questionNumber
    global op
    global opNum
    for id in selectQuestionlist:
        question = getQuestions(id)  # 获取问题
        optionsList = getOptions(id)  # 获取选项列表
        subject = document.add_paragraph(f'{questionNumber}、{question[0]}')  # 插入题目，questionNumber是题号
        subject.paragraph_format.line_spacing = Pt(22)  # 题目的行距
        optionParagraph = document.add_paragraph('')  # 插入选项段落
        optionParagraph.paragraph_format.line_spacing = Pt(22)  # 选项的行距
        for optionNum, option in zip(op, optionsList):  # 循环插入选项
            optionParagraph.add_run(f"{optionNum}、{option}   ")
        questionNumber += 1


def addQuestion(questionList):
    global questionNumber
    for id in questionList:
        question = getQuestions(id)  # 获取问题
        subject = document.add_paragraph(f'{questionNumber}、{question[0]}')  # 插入题目，questionNumber是题号
        subject.paragraph_format.line_spacing = Pt(22)  # 题目的行距
        questionNumber += 1


def questionsType():
    pass


def strToList(str):
    numList = str.strip('[').strip(']').split(',')  # 将获取的字符串转化为list
    return numList


def grouping(topicList):
    for i in topicList:
        sql = f"select type from tb_questions where id = {i}"
        cursor.execute(sql)
        type = cursor.fetchall()[0][0]
        print(type)
        if type == '选择题':
            selectList.append(i)
        elif type == '填空题':
            gapArray.append(i)
        elif type == '简答题':
            answerArray.append(i)


if __name__ == '__main__':
    db = pymysql.connect(host="gz-cynosdbmysql-grp-84tdgg43.sql.tencentcdb.com", user="root", password="gdupt@3c807",
                         database="db", port=25120)
    cursor = db.cursor()  # 创建游标对象操作数据库

    op = ['A', 'B', 'C', 'D']  # 选择题的选项编号
    opNum = 0  # 用来给选择题选项循环赋予选项
    questionNumber = 1  # 整张试卷里所有小题的编号
    savePath = "D:\work\javaidea\Project\problemSystem\\src\main\\resources\\static\\pyCreateFile\\"  # 保存试卷的地址
    picturePath = "D:\work\javaidea\Project\problemSystem\\src\main\\resources\\static\\py\\得分.bmp"  # 得分图片的地址

    inputStr = sys.argv[1]  # 从控制台获取java传参
    # inputStr = input('sdasd:')
    topicList = strToList(inputStr)  # 将str转为list

    document = Document('D:\work\javaidea\Project\problemSystem\src\main\\resources\static\py\学校模板.docx')  # 打开模板文档
    document.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    document.styles['Normal'].font.size = Pt(12)  # 全局字号
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)  # 全局字色
    selectList = []  # 选择题列表
    gapArray = []  # 填空题列表
    answerArray = []  # 简答题列表
    grouping(topicList)

    if selectList:
        point1 = document.add_picture(picturePath)  # 大题开头加入得分图片
        p1 = document.add_paragraph('')  # 建立空段落，才能添加粗体文字
        p1.add_run('一、单项选择题（在括号里填入正确答案，每小题 2 分，共30 分）').bold = True  # p1段落都添加粗体文字
        addSelectQuestion(selectList)  # 调用函数添加选择题
    if gapArray:
        point2 = document.add_picture(picturePath)
        p2 = document.add_paragraph('')
        p2.add_run('二、填空题（在横线上填上正确答案，每空 2 分，共 20 分）').bold = True
        addQuestion(gapArray)
    if answerArray:
        point3 = document.add_picture(picturePath)
        p3 = document.add_paragraph('')
        p3.add_run('三、将下列逻辑函数化为最简与或式：（方法不限，写出化简步骤，每小题 5 分，共 15 分。）').bold = True
        addQuestion(answerArray)

    document.save(savePath + "test.docx")  # 保存试卷
    cursor.close()  # 关闭游标对象
    db.close()  # 关闭数据库
    print(topicList)
